#!/usr/bin/env python3
"""
doc_engine.py — 文档处理引擎 (Step 1-5 优化版)
==============================================
优化点：
1. 入库强制分片 (700字窗口)
2. 检索结果强制截断 (500字)
3. 检索契约提示 (Protocol Guidance)
4. 精准 doc_id 过滤
5. 安全摘要锚点 (不暴露原文)
"""

import sys
import os
import re
import json
import shutil
import hashlib
from pathlib import Path
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from core_engine import HermesProjectCore

class DocumentEngine(HermesProjectCore):
    def __init__(self):
        super().__init__()
        self.archive_dir = Path(
            self.config.get(
                "HERMES_PROJECT_DOCUMENT_ARCHIVE_DIR",
                str(self.project_root / "archive" / "documents"),
            )
        )

    def _chunk_text(self, text: str, size: int = 700, overlap: int = 100) -> list[str]:
        """Step 1: 强制分片逻辑"""
        if not text: return []
        chunks = []
        start = 0
        while start < len(text):
            end = start + size
            chunks.append(text[start:end])
            start += (size - overlap)
        return chunks

    def extract_text(self, path: str) -> str:
        ext = os.path.splitext(path)[1].lower()
        if ext == ".pdf":
            import pdf_inspector
            return pdf_inspector.extract_text(path)
        elif ext == ".docx":
            return self._extract_docx(path)
        elif ext in (".txt", ".md", ".log", ".csv"):
            for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
                try:
                    with open(path, "r", encoding=enc, errors="strict") as f:
                        return f.read()
                except: continue
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        elif ext in (".xlsx", ".xls"):
            return self._extract_excel(path)
        return ""

    def _extract_docx(self, path: str) -> str:
        import docx
        doc = docx.Document(path)
        chunks = []
        for p in doc.paragraphs:
            if p.text.strip(): chunks.append(p.text)
        for tbl_idx, table in enumerate(doc.tables):
            chunks.append(f"\n[表格 {tbl_idx + 1}]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                chunks.append("\t".join(cells))
        return "\n".join(chunks)

    def _extract_excel(self, path: str) -> str:
        from openpyxl import load_workbook
        wb = load_workbook(path, data_only=True)
        chunks = []
        for ws in wb.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            rows = [r for r in rows if any(v is not None for v in r)]
            if not rows: continue
            chunks.append(f"\n## Sheet: {ws.title}")
            for i, row in enumerate(rows):
                cells = ["" if v is None else str(v) for v in row]
                chunks.append(" | ".join(cells))
                if i == 0: chunks.append(" | ".join(["---"] * len(cells)))
        return "\n".join(chunks)

    def get_summary_anchor(self, path: str, limit: int = 500) -> dict:
        """Step 3: 安全摘要，不暴露原文前 500 字"""
        if not os.path.isfile(path): return {"error": "File not found"}
        sql = "SELECT id, title, summary, tags, category FROM documents WHERE file_path = ?"
        res = self.query_d1(sql, [path])
        if res.get("success") and res.get("result"):
            rows = res["result"][0].get("results", [])
            if rows: return {**rows[0], "status": "indexed"}
        
        # Raw fallback: 只返回统计信息
        with open(path, "rb") as f:
            md5_tag = hashlib.md5(f.read()).hexdigest()[:8]
        size = os.path.getsize(path)
        return {
            "title": os.path.basename(path),
            "summary": f"Unindexed local file ({size} bytes). Use 'doc' command to process.",
            "md5": md5_tag,
            "status": "raw",
            "file_size": size,
        }

    def search_all(self, query: str, limit: int = 3, content_chars: int = 500) -> dict:
        """Step 2, 4, 5: 检索逻辑优化"""
        results = {
            "metadata": [], 
            "semantic_chunks": [],
            "protocol_guidance": "RETRIEVAL_ONLY: Chunks are fragments. Do NOT verbatim copy. Cite DocID and summarize."
        }
        
        # 1. D1 元数据
        sql = "SELECT id, title, company, category, tags FROM documents WHERE title LIKE ? OR tags LIKE ? OR company LIKE ? LIMIT ?"
        res = self.query_d1(sql, [f"%{query}%", f"%{query}%", f"%{query}%", limit])
        
        if res.get("success") and res.get("result"):
            rows = res["result"][0].get("results", [])
            results["metadata"] = rows
            if rows:
                # Step 4: 优先 doc_id 过滤检索
                doc_ids = [str(r["id"]) for r in rows]
                chroma_res = self._query_chroma_filtered(query, doc_ids, limit)
                results["semantic_chunks"] = self._trim_chunks(chroma_res, content_chars)
                results["retrieval_mode"] = "doc_id_filtered"
                return results

        # 2. 全库 fallback
        chroma_res = self.query_chroma("documents", query, n_results=limit)
        results["semantic_chunks"] = self._trim_chunks(chroma_res, content_chars)
        results["retrieval_mode"] = "full_library_fallback"
        return results

    def _query_chroma_filtered(self, query: str, doc_ids: list[str], n: int) -> dict:
        try:
            import chromadb
            client = chromadb.HttpClient(host=self.chroma_host, port=self.chroma_port)
            col = client.get_collection(name="documents")
            # Step 4: $in 过滤逻辑
            return col.query(
                query_texts=[query],
                n_results=n,
                where={"doc_id": {"$in": doc_ids}}
            )
        except: return {}

    def _trim_chunks(self, res: dict, limit: int) -> list[dict]:
        """Step 2: 强制截断"""
        docs = res.get("documents", [[]])[0]
        metas = res.get("metadatas", [[]])[0]
        out = []
        for i in range(len(docs)):
            txt = str(docs[i])
            trimmed = txt if len(txt) <= limit else txt[:limit] + "..."
            out.append({"content": trimmed, "metadata": metas[i] if metas else {}})
        return out

    def _resolve_d1_id(self, file_path: str) -> int | None:
        """根据 file_path 查询 D1 中文档的 id"""
        res = self.query_d1("SELECT id FROM documents WHERE file_path = ?", [file_path])
        if res.get("success") and res.get("result"):
            rows = res["result"][0].get("results", [])
            if rows:
                return rows[0]["id"]
        return None

    def process(self, path: str, data: dict):
        if not os.path.isfile(path): return None
        with open(path, "rb") as f:
            src_md5 = hashlib.md5(f.read()).hexdigest()[:8]
        date_str = datetime.now().strftime("%Y%m%d")
        safe_company = re.sub(r'[\\/*?:"<>|]', "", data.get("company", "Unknown"))[:10]
        safe_title = re.sub(r'[\\/*?:"<>|]', "", data.get("title", "Doc"))[:15]
        new_name = f"{date_str}_{safe_company}_{safe_title}_{src_md5}{os.path.splitext(path)[1]}"
        dest_path = self.archive_dir / new_name

        self.archive_dir.mkdir(parents=True, exist_ok=True)
        if not dest_path.exists(): shutil.copy(path, dest_path)

        raw_text = data.get("raw_text") or self.extract_text(path)

        sql = """
        INSERT INTO documents (title, company, category, summary, tags, file_path)
        VALUES (?, ?, ?, ?, ?, ?)
        ON CONFLICT(file_path) DO UPDATE SET title=excluded.title, summary=excluded.summary, tags=excluded.tags
        """
        params = [data.get("title"), data.get("company"), data.get("category"), data.get("summary"), data.get("tags"), str(dest_path)]
        res = self.query_d1(sql, params)

        if res.get("success"):
            # 获取 D1 文档 id（新插入或已存在）
            d1_id = self._resolve_d1_id(str(dest_path))
            doc_id = str(d1_id) if d1_id else data.get("title", new_name)

            # Step 1: 批量分片入库 — document_id 使用 D1 id
            chunks = self._chunk_text(raw_text)
            meta = {k:v for k,v in data.items() if k!="raw_text" and isinstance(v,(str,int,float,bool))}
            self.sync_to_chroma_chunks("documents", doc_id, chunks, meta)
            return dest_path
        return None

if __name__ == "__main__":
    engine = DocumentEngine()
    if len(sys.argv) > 2:
        engine.process(sys.argv[1], json.loads(sys.argv[2]))
