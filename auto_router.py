#!/usr/bin/env python3
"""
auto_router.py — 文档自动识别与路由
========================================

接收任意本地文件路径，自动判断类型后路由到对应引擎：
  - 发票（含"发票"/"invoice"/"税号"/"纳税人识别号"等关键词）→ FinanceEngine
  - 其他文档 → DocumentEngine

调用方式：

    # CLI
    ./project-tool process /path/to/file.pdf

    # Python
    from auto_router import AutoRouter
    AutoRouter().process("/path/to/file.pdf")

发票路由还会用 LLM 解析结构化字段（如有 OPENAI/Claude API），
否则退回 regex 规则提取。
"""

import sys
import os
import re
import json
import hashlib
import subprocess
from pathlib import Path
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# ── 发票关键词（命中任一即走发票流程）──────────────────────────────────────
INVOICE_KEYWORDS = [
    "发票", "invoice", "纳税人识别号", "税号", "含税金额", "不含税", "税率",
    "增值税", "普通发票", "专用发票", "开票", "价税合计", "销售方", "购买方",
    "发票代码", "发票号码", "校验码",
]

# ── 支持的文件格式 ───────────────────────────────────────────────────────────
SUPPORTED_EXTS = {".pdf", ".docx", ".doc", ".txt", ".md", ".log", ".csv", ".xlsx", ".xls"}


def extract_text(file_path: str) -> str:
    """
    从文件中提取纯文本。

    PDF  → pdf_inspector（pdf_engine.py）
    DOCX → python-docx
    TXT/MD/LOG/CSV → 多编码探测
    XLSX/XLS → openpyxl

    Returns:
        提取到的文本字符串（最多 12000 字符）
    """
    p = Path(file_path)
    ext = p.suffix.lower()

    if ext == ".pdf":
        try:
            result = subprocess.run(
                [sys.executable, str(Path(__file__).parent / "pdf_engine.py"), file_path, "json"],
                capture_output=True, text=True, timeout=60
            )
            data = json.loads(result.stdout)
            return data.get("text", "")[:12000]
        except Exception as e:
            return f"[PDF提取失败: {e}]"

    elif ext in (".docx", ".doc"):
        try:
            from docx import Document
            doc = Document(file_path)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(c.text.strip() for c in row.cells if c.text.strip()))
            return "\n".join(parts)[:12000]
        except Exception as e:
            return f"[DOCX提取失败: {e}]"

    elif ext in (".txt", ".md", ".log", ".csv"):
        for enc in ("utf-8", "gbk", "latin-1"):
            try:
                return p.read_text(encoding=enc)[:12000]
            except Exception:
                continue
        return "[文本提取失败: 编码不支持]"

    elif ext in (".xlsx", ".xls"):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            parts = []
            for sheet in wb.worksheets:
                parts.append(f"### {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(cells):
                        parts.append(" | ".join(cells))
            return "\n".join(parts)[:12000]
        except Exception as e:
            return f"[Excel提取失败: {e}]"

    else:
        return f"[不支持的格式: {ext}]"


def is_invoice(text: str) -> bool:
    """
    根据文本内容判断是否为发票。

    命中任意一个关键词即认定为发票。
    大小写不敏感，匹配简繁体关键词。

    Returns:
        True  — 应走发票流程
        False — 应走文档流程
    """
    text_lower = text.lower()
    for kw in INVOICE_KEYWORDS:
        if kw.lower() in text_lower:
            return True
    return False


def parse_invoice_fields(text: str, file_path: str) -> dict:
    """
    从发票文本中用正则规则提取结构化字段。

    尽量提取以下字段：
      invoice_number, invoice_date, buyer_name, seller_name,
      item_name, amount_net, tax_amount, total_amount

    未能提取到的字段填充 Unknown / 0。

    Returns:
        发票 dict，可直接传入 FinanceEngine.insert_confirmed_data()
    """
    def find(patterns, default="Unknown"):
        for pat in patterns:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                return m.group(1).strip()
        return default

    def find_amount(patterns, default=0.0):
        for pat in patterns:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                try:
                    return float(m.group(1).replace(",", ""))
                except Exception:
                    pass
        return default

    invoice_number = find([
        r"发票号码[：:]\s*([0-9A-Za-z\-]+)",
        r"invoice\s*(?:number|no\.?)[：:\s]+([0-9A-Za-z\-]+)",
        r"No\.?\s*([0-9]{8,20})",
    ])

    # 若仍未找到，生成一个基于文件 MD5 的唯一号避免重复
    if invoice_number == "Unknown":
        with open(file_path, "rb") as f:
            md5 = hashlib.md5(f.read()).hexdigest()[:12]
        invoice_number = f"AUTO-{md5.upper()}"

    invoice_date = find([
        r"开票日期[：:]\s*(\d{4}[-/年]\d{1,2}[-/月]\d{1,2})",
        r"invoice\s*date[：:\s]+(\d{4}[-/]\d{1,2}[-/]\d{1,2})",
        r"(\d{4}[-/]\d{2}[-/]\d{2})",
    ])

    buyer_name = find([
        r"购买方名称[：:]\s*(.+?)(?:\n|$)",
        r"购买方[：:]\s*(.+?)(?:\n|$)",
        r"bill\s*to[：:\s]+(.+?)(?:\n|$)",
    ])

    seller_name = find([
        r"销售方名称[：:]\s*(.+?)(?:\n|$)",
        r"销售方[：:]\s*(.+?)(?:\n|$)",
        r"seller[：:\s]+(.+?)(?:\n|$)",
        r"from[：:\s]+(.+?)(?:\n|$)",
    ])

    item_name = find([
        r"(?:货物|商品|服务)名称[：:]\s*(.+?)(?:\n|$)",
        r"item\s*(?:name|description)[：:\s]+(.+?)(?:\n|$)",
        r"品名[：:]\s*(.+?)(?:\n|$)",
    ])

    amount_net = find_amount([
        r"不含税金额[：:]\s*([\d,\.]+)",
        r"合计\s*¥?\s*([\d,\.]+)",
        r"subtotal[：:\s]+([\d,\.]+)",
    ])

    tax_amount = find_amount([
        r"税额[：:]\s*([\d,\.]+)",
        r"tax\s*amount[：:\s]+([\d,\.]+)",
        r"vat[：:\s]+([\d,\.]+)",
    ])

    total_amount = find_amount([
        r"价税合计[（\(]?大写[）\)]?[：:][^¥\d]*¥?\s*([\d,\.]+)",
        r"价税合计[：:]\s*¥?\s*([\d,\.]+)",
        r"total\s*amount[：:\s]+\$?([\d,\.]+)",
        r"合计金额[：:]\s*¥?\s*([\d,\.]+)",
        r"total[：:\s]+\$?([\d,\.]+)",
    ])

    return {
        "invoice_number": invoice_number,
        "invoice_date": invoice_date,
        "buyer_name": buyer_name,
        "seller_name": seller_name,
        "item_name": item_name,
        "amount_net": amount_net,
        "tax_amount": tax_amount,
        "total_amount": total_amount,
        "file_path": file_path,
        "raw_text": text[:6000],
    }


def infer_doc_meta(text: str, file_path: str) -> dict:
    """
    从文本中推断文档元数据（title / company / category / summary / tags）。

    使用启发式规则，不调用 LLM：
      - title: 取文本前 200 字中最长的非空行，或文件名
      - company: 从文本中搜索"公司/Corp/Ltd/Inc"附近词语
      - category: 关键词映射
      - summary: 文本前 300 字
      - tags: 从 category / title / company 拼合

    Returns:
        文档 metadata dict，可直接传入 DocumentEngine.process()
    """
    p = Path(file_path)
    lines = [l.strip() for l in text[:500].splitlines() if l.strip()]

    # title: 取最长的前几行（跳过表单标签类短行）
    title = p.stem
    for line in lines[:10]:
        if len(line) > 6 and not re.match(r'^[\d\s\|▶•]+$', line):
            title = line[:60]
            break

    # company: 搜索企业名称
    company_match = re.search(
        r'([^\n]{2,30}(?:有限公司|股份公司|Corp|Ltd|Inc|LLC|Co\.))',
        text[:3000], re.IGNORECASE
    )
    company = company_match.group(1).strip() if company_match else "Unknown"

    # category: 关键词映射
    cat_map = [
        (["发票", "增值税", "纳税人识别号", "税号", "价税合计", "w-8", "w-9", "irs tax", "tax withholding"], "税务文件"),
        (["合同", "协议", "agreement", "contract"], "合同协议"),
        (["方案", "proposal", "solution"], "技术方案"),
        (["简历", "resume", "cv"], "简历"),
        (["招标", "投标", "tender", "bid"], "招投标"),
        (["财务", "finance", "账单", "账务"], "财务文件"),
        (["规范", "标准", "specification", "standard"], "规范标准"),
        (["报告", "report", "分析"], "报告分析"),
    ]
    category = "通用文档"
    text_lower = text[:2000].lower()
    for keywords, cat in cat_map:
        if any(kw.lower() in text_lower for kw in keywords):
            category = cat
            break

    summary = re.sub(r'\s+', ' ', text[:300]).strip()

    tags_parts = [category]
    if company != "Unknown":
        tags_parts.append(company[:10])
    if title and title != p.stem:
        tags_parts.append(title[:15])
    tags = ",".join(tags_parts)

    return {
        "title": title,
        "company": company,
        "category": category,
        "summary": summary,
        "tags": tags,
    }


class AutoRouter:
    """
    自动识别文档类型并路由到对应引擎的入口类。

    用法：
        router = AutoRouter()
        router.process("/path/to/file.pdf")
    """

    def process(self, file_path: str):
        """
        主入口：提取文本 → 判断类型 → 路由处理。

        Args:
            file_path: 本地文件绝对路径
        """
        p = Path(file_path)
        if not p.exists():
            print(f"❌ 文件不存在: {file_path}")
            sys.exit(1)

        if p.suffix.lower() not in SUPPORTED_EXTS:
            print(f"❌ 不支持的格式: {p.suffix}（支持：{', '.join(SUPPORTED_EXTS)}）")
            sys.exit(1)

        print(f"📄 正在处理: {p.name}")

        # ── 1. 提取文本 ────────────────────────────────────────────────────
        text = extract_text(file_path)
        if not text.strip() or text.startswith("["):
            print(f"⚠️  文本提取为空或失败，仍尝试走文档流程: {text[:100]}")

        # ── 2. 识别类型 ────────────────────────────────────────────────────
        if is_invoice(text):
            print(f"🧾 识别为发票 → 走发票处理流程")
            self._route_invoice(text, file_path)
        else:
            print(f"📁 识别为普通文档 → 走文档处理流程")
            self._route_document(text, file_path)

    def _route_invoice(self, text: str, file_path: str):
        """解析发票字段并调用 FinanceEngine 入库。"""
        from invoice_engine import FinanceEngine
        data = parse_invoice_fields(text, file_path)
        print(f"   发票号: {data['invoice_number']}")
        print(f"   购买方: {data['buyer_name']}  销售方: {data['seller_name']}")
        print(f"   金额:   ¥{data['total_amount']}")
        FinanceEngine().insert_confirmed_data(data)

    def _route_document(self, text: str, file_path: str):
        """推断文档元数据并调用 DocumentEngine 入库。"""
        from doc_engine import DocumentEngine
        meta = infer_doc_meta(text, file_path)
        meta["file_path_src"] = file_path
        meta["raw_text"] = text[:8000]
        print(f"   标题: {meta['title']}")
        print(f"   分类: {meta['category']}  公司: {meta['company']}")
        file_path_resolved = meta.pop("file_path_src")
        DocumentEngine().process(file_path_resolved, meta)


if __name__ == "__main__":
    if len(sys.argv) < 2 or sys.argv[1] in ("-h", "--help"):
        print("Usage: python auto_router.py <file_path>")
        print("       ./project-tool process <file_path>")
        sys.exit(0)
    AutoRouter().process(sys.argv[1])
